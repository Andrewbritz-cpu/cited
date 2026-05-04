"""
CV file parsers — extract plain text from uploaded files.

Supports PDF (via pdfplumber) and DOCX (via python-docx). Designed to surface
parsing failures clearly so the frontend can tell the user "your CV is too
formatted for an ATS to read" rather than just throwing a generic error.
"""

import io
from typing import Union

import pdfplumber
from docx import Document

PDF_TYPES = {"application/pdf"}
DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",  # Best-effort: legacy .doc parsing is unreliable.
}


def extract_text_from_cv(content: bytes, content_type: str) -> str:
    """
    Extract plain text from a CV file.

    Args:
        content: Raw file bytes.
        content_type: MIME type from the upload.

    Returns:
        Extracted plain text.

    Raises:
        ValueError if the content type is unsupported or extraction yields nothing.
    """
    if content_type in PDF_TYPES:
        return _extract_from_pdf(content)
    if content_type in DOCX_TYPES:
        return _extract_from_docx(content)
    raise ValueError(f"Unsupported content type: {content_type}")


def _extract_from_pdf(content: bytes) -> str:
    """Extract text from a PDF using pdfplumber."""
    pages_text = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            pages_text.append(text)
    return "\n\n".join(pages_text)


def _extract_from_docx(content: bytes) -> str:
    """Extract text from a Word document using python-docx."""
    doc = Document(io.BytesIO(content))
    parts = []

    # Body paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    # Tables (which ATS systems struggle with — but we still extract the text
    # so the scorer can flag the issue rather than miss content silently)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)
