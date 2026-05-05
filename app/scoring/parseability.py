"""
Parseability scoring (40% of total score).

This is the most heavily weighted component because it's the single biggest
real-world reason CVs fail ATS systems: the parser literally can't read them.
We use the *actual* output of pdfplumber/python-docx to detect parsing
failures, because real ATS parsers fail on the same things.

Inputs we work with:
- raw_text: the text we successfully extracted from the file
- raw_bytes: the original file bytes (so we can re-parse to detect specific
  failure modes without re-uploading)
- content_type: PDF or DOCX

Output: parseability score (0-40) and a list of specific issues.
"""

import io
import re
from typing import List, Tuple

import pdfplumber
from docx import Document

from app.models import StructuralIssue


def score_parseability(
    raw_text: str,
    raw_bytes: bytes,
    content_type: str,
) -> Tuple[int, List[StructuralIssue]]:
    """
    Score parseability out of 40 points and return the list of detected issues.

    Issues are ordered by severity (critical first). The score is calculated
    by starting at 40 and subtracting per-issue penalties, capped at 0.
    """
    issues: List[StructuralIssue] = []

    # ---- A: text yield is suspiciously low ----
    # Real CVs typically have 1500-5000 characters of body text. Below 500 is
    # a strong signal of image-based content, embedded text boxes, or tables
    # that the parser couldn't navigate.
    char_count = len(raw_text.strip())
    if char_count < 200:
        issues.append(StructuralIssue(
            severity="critical",
            type="parser_extracted_almost_nothing",
            description=(
                "We extracted under 200 characters from this CV. That almost "
                "always means the content is inside images, text boxes, or "
                "complex tables that ATS systems can't read either."
            ),
            penalty=35,
        ))
    elif char_count < 500:
        issues.append(StructuralIssue(
            severity="high",
            type="low_text_extraction",
            description=(
                "We extracted only a small amount of text. ATS systems rely on "
                "text extraction; if we struggle, they will too. This usually "
                "indicates heavy use of graphics, tables, or text boxes."
            ),
            penalty=20,
        ))

    # ---- B: format-specific structural issues ----
    if content_type == "application/pdf":
        pdf_issues = _detect_pdf_issues(raw_bytes, raw_text)
        issues.extend(pdf_issues)
    elif content_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }:
        docx_issues = _detect_docx_issues(raw_bytes, raw_text)
        issues.extend(docx_issues)

    # ---- C: encoding artefacts ----
    # Mojibake or replacement characters indicate the file came from a
    # non-standard source (scanned, OCR'd, or exported from an unusual tool).
    replacement_count = raw_text.count("\ufffd") + raw_text.count("\x00")
    if replacement_count > 5:
        issues.append(StructuralIssue(
            severity="high",
            type="encoding_problems",
            description=(
                "Text encoding errors detected — multiple unreadable characters "
                "in the extracted content. ATS systems will store these as "
                "garbage, breaking keyword matching."
            ),
            penalty=12,
        ))

    # ---- D: line-length distribution suggests multi-column layout ----
    # When pdfplumber encounters multi-column layouts, it usually emits very
    # short lines (one column word) interleaved with another column's words.
    if char_count >= 200:
        issues.extend(_detect_multicolumn_via_line_lengths(raw_text))

    # ---- E: missing whitespace runs (one giant blob of text) ----
    # The opposite failure mode: text without paragraph breaks usually means
    # the parser collapsed structural whitespace.
    if char_count > 1000 and raw_text.count("\n") < 8:
        issues.append(StructuralIssue(
            severity="medium",
            type="lost_paragraph_structure",
            description=(
                "Text was extracted as one continuous block with very few line "
                "breaks. ATS systems use paragraph breaks to identify sections; "
                "without them, the CV reads as one undifferentiated lump."
            ),
            penalty=8,
        ))

    # ---- Final score ----
    total_penalty = sum(issue.penalty for issue in issues)
    score = max(40 - total_penalty, 0)

    # Sort issues by severity (critical first) so callers get the worst news first
    severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3}
    issues.sort(key=lambda i: severity_order.get(i.severity, 9))

    return score, issues


def _detect_pdf_issues(raw_bytes: bytes, raw_text: str) -> List[StructuralIssue]:
    """Re-parse the PDF to look at structure pdfplumber didn't surface as text."""
    issues: List[StructuralIssue] = []

    try:
        with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
            page_count = len(pdf.pages)
            total_chars = sum(len(p.chars) for p in pdf.pages)
            total_images = sum(len(p.images) for p in pdf.pages)
            total_tables = sum(len(p.find_tables()) for p in pdf.pages)

            # Heavy table use is a classic ATS-killer
            if total_tables >= 2:
                issues.append(StructuralIssue(
                    severity="high",
                    type="multiple_tables_detected",
                    description=(
                        f"Detected {total_tables} tables in the PDF. ATS parsers "
                        "frequently misread tables — they may transpose columns, "
                        "merge cells, or skip them entirely."
                    ),
                    penalty=15,
                ))

            # Image-heavy CV with little text = probably a designed CV
            if total_images >= 3 and total_chars < 1500:
                issues.append(StructuralIssue(
                    severity="high",
                    type="image_heavy_design",
                    description=(
                        f"Detected {total_images} images but only {total_chars} "
                        "characters of extractable text. Designed CVs look great "
                        "to humans but are largely invisible to ATS systems."
                    ),
                    penalty=18,
                ))

            # Page count outside conventions
            if page_count > 5:
                issues.append(StructuralIssue(
                    severity="medium",
                    type="excessive_page_count",
                    description=(
                        f"This CV is {page_count} pages long. Most ATS systems "
                        "and recruiters favour 2-3 page CVs; extreme length can "
                        "trigger filtering or simply lose the reader's attention."
                    ),
                    penalty=6,
                ))

    except Exception as exc:
        # PDF re-parsing failed entirely — that itself is the issue
        issues.append(StructuralIssue(
            severity="critical",
            type="pdf_structure_unreadable",
            description=(
                "The PDF structure couldn't be re-parsed cleanly. This often "
                "indicates a corrupted file, an image-only PDF (scan), or "
                "an unusual export format. Most ATS systems would fail here too."
            ),
            penalty=30,
        ))

    return issues


def _detect_docx_issues(raw_bytes: bytes, raw_text: str) -> List[StructuralIssue]:
    """Inspect DOCX structure for ATS-hostile features."""
    issues: List[StructuralIssue] = []

    try:
        doc = Document(io.BytesIO(raw_bytes))

        # Text inside tables — modern ATSs handle simple tables OK, but heavy
        # use is still risky
        table_count = len(doc.tables)
        if table_count >= 3:
            issues.append(StructuralIssue(
                severity="high",
                type="multiple_tables_in_docx",
                description=(
                    f"Detected {table_count} tables in the document. Tables "
                    "remain one of the most common reasons CVs get mis-parsed; "
                    "convert to plain text where possible."
                ),
                penalty=12,
            ))

        # Header/footer use — ATSs frequently skip them, losing contact info
        for section in doc.sections:
            if section.header.paragraphs:
                header_text = "\n".join(p.text for p in section.header.paragraphs).strip()
                if len(header_text) > 10:
                    # Check whether contact info is *only* in the header
                    if "@" in header_text and "@" not in raw_text:
                        issues.append(StructuralIssue(
                            severity="critical",
                            type="contact_only_in_header",
                            description=(
                                "Email address found only in the document header. "
                                "Many ATS systems skip headers entirely, meaning "
                                "your contact details may be invisible to recruiters."
                            ),
                            penalty=25,
                        ))
                    break  # Only need to flag header issue once

    except Exception:
        # If we can't even open it as DOCX, the parser already extracted what
        # it could into raw_text — no additional issue to flag.
        pass

    return issues


def _detect_multicolumn_via_line_lengths(raw_text: str) -> List[StructuralIssue]:
    """
    Multi-column layouts produce a distinctive line-length distribution:
    many very short lines (1-3 words) interleaved with full-width text.

    A clean single-column CV has lines that average 30-80 characters with
    relatively low variance. A multi-column extraction has bimodal lengths:
    very short (column-only) and very long (concatenated columns).
    """
    issues: List[StructuralIssue] = []

    lines = [line for line in raw_text.split("\n") if line.strip()]
    if len(lines) < 20:
        return issues  # Not enough data to make a judgement

    short_lines = sum(1 for line in lines if 0 < len(line.strip()) < 15)
    short_ratio = short_lines / len(lines)

    if short_ratio > 0.45:
        issues.append(StructuralIssue(
            severity="high",
            type="multicolumn_layout_suspected",
            description=(
                "Line-length distribution suggests a multi-column layout was "
                "extracted into jumbled order. ATS parsers usually concatenate "
                "columns left-to-right rather than reading each column "
                "top-to-bottom, scrambling your CV's meaning."
            ),
            penalty=14,
        ))

    return issues
